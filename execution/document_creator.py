"""
Stewie Document Creator — Generate Word documents from research and content.

Creates professionally formatted .docx files with titles, headings,
bullet points, and source citations.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional, Union

from loguru import logger


async def create_document(
    title: str,
    content: Union[str, dict],
    filename: Optional[str] = None,
    save_path: Optional[str] = None,
) -> str:
    """
    Create a Microsoft Word document with the given content.

    Args:
        title: Document title.
        content: Text content or research results dict.
        filename: Desired filename (without extension).
        save_path: Directory to save in. Defaults to Desktop.

    Returns:
        Path to the saved document.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise

    doc = Document()

    # --- Document Styling ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # --- Title ---
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Metadata ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Compiled by Stewie AI Assistant — "
        f"{datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")  # Spacer

    # --- Content ---
    if isinstance(content, dict):
        # Research results format
        _populate_research_doc(doc, content)
    elif isinstance(content, str):
        # Plain text — parse markdown-like formatting
        _populate_text_doc(doc, content)
    else:
        doc.add_paragraph(str(content))

    # --- Save ---
    if not filename:
        # Generate filename from title
        filename = re.sub(r'[<>:"/\\|?*]', "", title)
        filename = filename.replace(" ", " ")[:80]

    if not save_path:
        from config.settings import load_config

        config = load_config()
        save_path = str(config.resolved_save_path)

    save_dir = Path(save_path)
    save_dir.mkdir(parents=True, exist_ok=True)

    full_path = save_dir / f"{filename}.docx"

    # Handle duplicate filenames
    counter = 1
    while full_path.exists():
        full_path = save_dir / f"{filename} ({counter}).docx"
        counter += 1

    doc.save(str(full_path))
    logger.info(f"Document saved: {full_path}")

    return str(full_path)


def _populate_research_doc(doc, research: dict) -> None:
    """Populate a Word doc from research results dict."""
    from docx.shared import Pt

    # Summary section
    if research.get("summary"):
        doc.add_heading("Summary", level=1)
        _add_formatted_paragraphs(doc, research["summary"])

    # Key points section
    key_points = research.get("key_points", [])
    if key_points:
        doc.add_heading("Key Findings", level=1)
        for point in key_points:
            doc.add_paragraph(point, style="List Bullet")

    # Sources section
    sources = research.get("sources", [])
    if sources:
        doc.add_heading("Sources", level=1)
        for i, source in enumerate(sources, 1):
            if isinstance(source, dict):
                title = source.get("title", f"Source {i}")
                url = source.get("url", "")
                para = doc.add_paragraph(style="List Number")
                run = para.add_run(title)
                run.bold = True
                if url:
                    para.add_run(f"\n{url}").font.size = Pt(9)
            else:
                doc.add_paragraph(str(source), style="List Number")


def _populate_text_doc(doc, text: str) -> None:
    """Populate a Word doc from plain/markdown-like text."""
    lines = text.split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        # Heading detection (markdown-style)
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith(("- ", "• ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        # Numbered lists
        elif re.match(r"^\d+\.\s", stripped):
            content = re.sub(r"^\d+\.\s*", "", stripped)
            doc.add_paragraph(content, style="List Number")
        else:
            doc.add_paragraph(stripped)


def _add_formatted_paragraphs(doc, text: str) -> None:
    """Add text as paragraphs, splitting on double newlines."""
    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        clean = para_text.strip()
        if clean:
            _populate_text_doc(doc, clean)
